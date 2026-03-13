"""
PrivateProxy — DOCX file parser.
Extracts text content from Microsoft Word documents.
"""

from __future__ import annotations

import logging
from pathlib import Path

from docx import Document

logger = logging.getLogger("privateproxy.parsers.docx")


def parse_docx(file_path: Path) -> str:
    """
    Extract all text from a DOCX file.

    Extracts text from:
    - Paragraphs (body text)
    - Tables (cell content)
    - Headers and footers

    Args:
        file_path: Path to the .docx file.

    Returns:
        Full text content of the document as a single string.
    """
    doc = Document(str(file_path))
    parts: list[str] = []

    # Extract paragraphs
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    # Extract table contents
    for table in doc.tables:
        for row in table.rows:
            row_texts = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_texts.append(cell_text)
            if row_texts:
                parts.append(" | ".join(row_texts))

    # Extract headers
    for section in doc.sections:
        header = section.header
        if header:
            for paragraph in header.paragraphs:
                text = paragraph.text.strip()
                if text:
                    parts.append(f"[Nagłówek] {text}")

        footer = section.footer
        if footer:
            for paragraph in footer.paragraphs:
                text = paragraph.text.strip()
                if text:
                    parts.append(f"[Stopka] {text}")

    full_text = "\n".join(parts)
    logger.info(f"Parsed DOCX: {file_path.name} — {len(parts)} text blocks, {len(full_text)} chars")
    return full_text
