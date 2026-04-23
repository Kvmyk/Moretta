"""
Moretta — Test configuration and fixtures.
Uses FastAPI TestClient with SSO disabled for isolated testing.
"""

import os
import sys
import tempfile
from pathlib import Path

import pytest

# Ensure backend is importable
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

# Override settings BEFORE importing the app
os.environ.update({
    "SSO_ENABLED": "false",
    "DATA_DIR": tempfile.mkdtemp(prefix="moretta_test_"),
    "DATABASE_BACKEND": "sqlite",
    "OLLAMA_URL": "http://localhost:11434",
    "LOCAL_MODEL": "phi4-mini",
    "LOG_LEVEL": "WARNING",
    "ANTHROPIC_API_KEY": "",
    "OPENAI_API_KEY": "",
    "GOOGLE_AI_API_KEY": "",
    "OPENROUTER_API_KEY": "",
})


@pytest.fixture(scope="session")
def test_data_dir():
    """Return the temporary data directory used for tests."""
    return Path(os.environ["DATA_DIR"])


@pytest.fixture()
def client():
    """Create a fresh TestClient for each test."""
    # Clear settings cache so env overrides take effect
    from config import get_settings
    get_settings.cache_clear()

    from fastapi.testclient import TestClient
    from main import app, file_store, task_store

    # Initialize stores for testing
    file_store.initialize()
    task_store.initialize()

    with TestClient(app) as c:
        yield c


@pytest.fixture()
def sample_docx(tmp_path: Path) -> Path:
    """Create a minimal .docx file for upload testing."""
    pytest.importorskip("docx")
    from docx import Document

    doc = Document()
    doc.add_paragraph("Jan Kowalski, PESEL 92010212345, mieszka w Warszawie.")
    doc.add_paragraph("Numer telefonu: 600 123 456, email: jan@firma.pl")
    path = tmp_path / "test_document.docx"
    doc.save(str(path))
    return path


@pytest.fixture()
def sample_xlsx(tmp_path: Path) -> Path:
    """Create a minimal .xlsx file for upload testing."""
    pytest.importorskip("openpyxl")
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Dane"
    ws["A1"] = "Imie i nazwisko"
    ws["B1"] = "PESEL"
    ws["A2"] = "Jan Kowalski"
    ws["B2"] = "92010212345"
    path = tmp_path / "test_sheet.xlsx"
    wb.save(str(path))
    return path


@pytest.fixture()
def sample_pdf(tmp_path: Path) -> Path:
    """Create a minimal .pdf file for upload testing."""
    pytest.importorskip("reportlab")
    from reportlab.pdfgen import canvas

    path = tmp_path / "test_document.pdf"
    pdf = canvas.Canvas(str(path))
    pdf.drawString(100, 750, "Jan Kowalski, PESEL 92010212345")
    pdf.drawString(100, 730, "Numer telefonu: 600 123 456")
    pdf.save()
    return path


@pytest.fixture()
def sample_text() -> str:
    """Sample text containing PII for testing."""
    return "Umowa pomiędzy Jan Kowalski (PESEL 92010212345) a firmą ABC Sp. z o.o."
