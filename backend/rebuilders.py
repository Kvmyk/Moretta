"""
Moretta — Rebuilders.
Reconstruct files (DOCX, XLSX) from AI-processed text,
preserving the original document's formatting/styles.
"""

import io
import os
import copy
import openpyxl
import logging
from docx import Document
from docx.oxml.ns import qn
from pathlib import Path
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

logger = logging.getLogger("moretta.rebuilders")


def rebuild_xlsx(text: str, template_path: str | None = None) -> bytes:
    """
    Reconstruct an XLSX file.
    If a template is provided, it preserves column widths, cell styles,
    merged cells, etc. — but replaces ALL cell values with AI output.
    """
    try:
        if template_path and os.path.exists(template_path):
            logger.info(f"Rebuilding XLSX using template: {template_path}")
            wb = openpyxl.load_workbook(template_path)

            # Step 1: Clear ALL cell values in the template (keep styles)
            for ws in wb.worksheets:
                for row in ws.iter_rows():
                    for cell in row:
                        if type(cell).__name__ != 'MergedCell':
                            cell.value = None
        else:
            logger.info("Rebuilding XLSX without template")
            wb = openpyxl.Workbook()
            if wb.active:
                wb.remove(wb.active)

        # Step 2: Fill in the AI text
        current_sheet = None
        row_offset = 0

        import re
        # Pattern to match cell coordinates like "A1: Value" or "AB123: Value"
        cell_pattern = re.compile(r"^([a-zA-Z]+)(\d+):\s*(.*)$")
        from openpyxl.utils import column_index_from_string

        max_row = 1

        for line in text.splitlines():
            line = line.strip()
            if not line:
                continue

            if line.startswith("[Arkusz:") and line.endswith("]"):
                sheet_name = line[8:-1].strip()
                if not sheet_name:
                    sheet_name = "Sheet1"

                if sheet_name in wb.sheetnames:
                    current_sheet = wb[sheet_name]
                    logger.debug(f"Writing to existing sheet: {sheet_name}")
                else:
                    current_sheet = wb.create_sheet(title=sheet_name[:31])
                    logger.debug(f"Created new sheet: {sheet_name}")
                row_offset = 0
                max_row = 1
                continue

            if current_sheet is None:
                current_sheet = wb.active or wb.create_sheet(title="Wynik")

            match = cell_pattern.match(line)
            if match:
                col_letter = match.group(1).upper()
                row_idx = int(match.group(2))
                val = match.group(3)
                
                max_row = max(max_row, row_idx)
                
                try:
                    col_idx = column_index_from_string(col_letter)
                    cell = current_sheet.cell(row=row_idx, column=col_idx)
                    if type(cell).__name__ != 'MergedCell':
                        cell.value = val
                except Exception as e:
                    logger.warning(f"Failed to write cell {col_letter}{row_idx}: {e}")
            else:
                # Fallback: if AI generated plain text, append it into the next available row in Column A
                max_row += 1
                try:
                    cell = current_sheet.cell(row=max_row, column=1)
                    if type(cell).__name__ != 'MergedCell':
                        cell.value = line
                except Exception as e:
                    pass

        stream = io.BytesIO()
        wb.save(stream)
        return stream.getvalue()
    except Exception as exc:
        logger.error(f"Error in rebuild_xlsx: {str(exc)[:200]}")
        raise


def rebuild_docx(text: str, template_path: str | None = None) -> bytes:
    """
    Reconstruct a DOCX file.
    If a template is provided, it preserves:
      - Page margins, headers, footers
      - Font styles from the first paragraph as a reference
      - Section layout
    All content is replaced with AI output — no leftover original text.
    """
    try:
        if template_path and os.path.exists(template_path):
            logger.info(f"Rebuilding DOCX using template: {template_path}")
            doc = Document(template_path)
            lines = [l for l in text.splitlines() if l.strip()]

            # Collect style info from existing paragraphs for reuse
            style_map = []
            for p in doc.paragraphs:
                style_info = {
                    "style": p.style,
                    "alignment": p.alignment,
                    "runs": [],
                }
                for run in p.runs:
                    style_info["runs"].append({
                        "bold": run.bold,
                        "italic": run.italic,
                        "underline": run.underline,
                        "font_name": run.font.name,
                        "font_size": run.font.size,
                        "font_color": run.font.color.rgb if run.font.color and run.font.color.rgb else None,
                    })
                style_map.append(style_info)

            # Step 1: Remove ALL existing paragraphs' content
            for p in doc.paragraphs:
                # Clear all runs
                for run in p.runs:
                    run.text = ""
                # Also clear any direct text
                p.text = ""

            # Step 2: Fill in AI text
            num_existing = len(doc.paragraphs)
            for i, line in enumerate(lines):
                if i < num_existing:
                    para = doc.paragraphs[i]
                    # Re-apply style from style_map if available
                    if i < len(style_map):
                        para.style = style_map[i]["style"]
                        para.alignment = style_map[i]["alignment"]

                    # If original had runs with formatting, use the first run's style
                    if i < len(style_map) and style_map[i]["runs"]:
                        # Clear paragraph XML children that are runs
                        for r_elem in para._element.findall(qn('w:r')):
                            para._element.remove(r_elem)
                        # Add a fresh run with old style
                        new_run = para.add_run(line.strip())
                        rs = style_map[i]["runs"][0]
                        new_run.bold = rs["bold"]
                        new_run.italic = rs["italic"]
                        new_run.underline = rs["underline"]
                        if rs["font_name"]:
                            new_run.font.name = rs["font_name"]
                        if rs["font_size"]:
                            new_run.font.size = rs["font_size"]
                        if rs["font_color"]:
                            new_run.font.color.rgb = rs["font_color"]
                    else:
                        para.text = line.strip()
                else:
                    # More AI lines than original paragraphs — add new ones
                    # Try to inherit style from last available paragraph
                    new_para = doc.add_paragraph(line.strip())
                    if style_map:
                        last_style = style_map[-1]
                        new_para.style = last_style["style"]
                        new_para.alignment = last_style["alignment"]

            logger.info(f"DOCX rebuilt: {len(lines)} AI lines into {num_existing} original paragraphs")
        else:
            logger.info("Rebuilding DOCX without template")
            doc = Document()
            for line in text.splitlines():
                if line.strip():
                    doc.add_paragraph(line.strip())

        stream = io.BytesIO()
        doc.save(stream)
        return stream.getvalue()
    except Exception as exc:
        logger.error(f"Error in rebuild_docx: {str(exc)[:200]}")
        raise


def rebuild_pdf(text: str) -> bytes:
    """
    Reconstruct a PDF file from plain text.
    """
    try:
        stream = io.BytesIO()
        pdf = canvas.Canvas(stream, pagesize=A4)
        width, height = A4

        margin_x = 50
        margin_y = 50
        line_height = 14
        y = height - margin_y

        for paragraph in text.splitlines():
            line = paragraph.rstrip()
            if not line:
                y -= line_height
                if y < margin_y:
                    pdf.showPage()
                    y = height - margin_y
                continue

            chunks = [line[i:i + 110] for i in range(0, len(line), 110)]
            for chunk in chunks:
                pdf.drawString(margin_x, y, chunk)
                y -= line_height
                if y < margin_y:
                    pdf.showPage()
                    y = height - margin_y

        pdf.save()
        return stream.getvalue()
    except Exception as exc:
        logger.error(f"Error in rebuild_pdf: {str(exc)[:200]}")
        raise
